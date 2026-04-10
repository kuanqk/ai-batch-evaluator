"""Extract text from docx/pdf: python-docx → PyMuPDF → Vision OCR (Qwen3-VL)."""

from __future__ import annotations

import base64
import io
import logging
import re
from pathlib import Path

import fitz
from django.conf import settings
from openai import AsyncOpenAI

from config.concurrency import vision_semaphore

from pipeline.converter import convert_docx_to_pdf
from pipeline.docx_utils import extract_text_from_docx_xml

logger = logging.getLogger(__name__)


def _min_chars_docx() -> int:
    return int(getattr(settings, "MIN_TEXT_CHARS", 50))


def _min_chars_pdf() -> int:
    # Slightly higher bar for born-digital PDF text layer
    return max(_min_chars_docx(), 100)


def _vision_max_pages() -> int:
    return int(getattr(settings, "VISION_MAX_PAGES", 10))


def _vision_dpi() -> int:
    return int(getattr(settings, "VISION_DPI", 150))


def _table_to_markdown(table) -> str:
    """Convert a python-docx table to GitHub-style markdown."""
    rows_out: list[str] = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            t = (cell.text or "").strip().replace("\n", " ").replace("|", "\\|")
            cells.append(t)
        rows_out.append("| " + " | ".join(cells) + " |")
    if not rows_out:
        return ""
    ncols = len(table.rows[0].cells)
    sep = "| " + " | ".join(["---"] * ncols) + " |"
    return rows_out[0] + "\n" + sep + "\n" + "\n".join(rows_out[1:])


def extract_text_with_python_docx(file_content: bytes) -> str:
    """Extract text from docx via python-docx; tables as markdown."""
    from docx import Document

    try:
        doc = Document(io.BytesIO(file_content))
    except Exception as e:
        logger.debug("python-docx open failed: %s", e)
        return ""
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        md = _table_to_markdown(table)
        if md:
            parts.append(md)
    return "\n\n".join(parts).strip()


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    text_parts: list[str] = []
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as e:
        logger.warning("PyMuPDF open failed: %s", e)
        return ""
    try:
        for i in range(len(doc)):
            page = doc.load_page(i)
            text_parts.append(page.get_text() or "")
    finally:
        doc.close()
    text = "\n".join(text_parts)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


async def extract_text_via_vision_qwen(
    pdf_content: bytes,
    *,
    max_pages: int | None = None,
    dpi: int | None = None,
) -> str:
    """PDF → PNG per page (base64) → NITEC vision model → concatenated text."""
    max_pages = max_pages if max_pages is not None else _vision_max_pages()
    dpi = dpi if dpi is not None else _vision_dpi()

    sem = vision_semaphore
    if sem is None:
        raise RuntimeError("vision_semaphore not initialized; call init_concurrency() first")
    if not settings.NITEC_API_KEY:
        raise ValueError("NITEC_API_KEY is not set for Vision OCR")

    try:
        doc = fitz.open(stream=pdf_content, filetype="pdf")
    except Exception as e:
        logger.warning("vision OCR: cannot open PDF: %s", e)
        return ""

    images_b64: list[str] = []
    try:
        n = min(len(doc), max_pages)
        matrix = fitz.Matrix(dpi / 72.0, dpi / 72.0)
        for i in range(n):
            page = doc.load_page(i)
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            png = pix.tobytes("png")
            images_b64.append(base64.b64encode(png).decode("ascii"))
    finally:
        doc.close()

    if not images_b64:
        return ""

    client = AsyncOpenAI(
        base_url=settings.NITEC_BASE_URL.rstrip("/"),
        api_key=settings.NITEC_API_KEY,
    )

    content: list[dict] = [
        {
            "type": "text",
            "text": (
                "Извлеки весь читаемый текст с этих страниц документа. "
                "Сохрани порядок и структуру абзацев. Без комментариев, только текст."
            ),
        }
    ]
    for b64 in images_b64:
        content.append(
            {
                "type": "image_url",
                "image_url": {"url": f"data:image/png;base64,{b64}"},
            }
        )

    async with sem:
        response = await client.chat.completions.create(
            model=settings.NITEC_VISION_MODEL,
            temperature=0.0,
            max_tokens=8192,
            messages=[{"role": "user", "content": content}],
        )

    raw = response.choices[0].message.content
    return (raw or "").strip()


async def extract_text(content: bytes, filename: str) -> tuple[str, str]:
    """
    Returns (text, method) where method is one of:
    plain | python_docx | xml | pdf_text | vision_ocr | empty
    """
    min_d = _min_chars_docx()
    min_pdf = _min_chars_pdf()

    ext = Path(filename).suffix.lower()

    if ext == ".txt":
        try:
            t = content.decode("utf-8")
        except UnicodeDecodeError:
            t = content.decode("utf-8", errors="replace")
        t = t.strip()
        return t, "plain"

    if ext == ".pdf":
        text = extract_text_from_pdf(content)
        if len(text) >= min_pdf:
            return text, "pdf_text"
        if len(text) < min_pdf:
            ocr = await extract_text_via_vision_qwen(content)
            if len(ocr.strip()) >= min_pdf:
                return ocr.strip(), "vision_ocr"
        return text, "pdf_text" if text else "empty"

    if ext == ".docx":
        text = extract_text_with_python_docx(content)
        if len(text) >= min_d:
            return text, "python_docx"
        text = extract_text_from_docx_xml(content)
        if len(text) >= min_d:
            return text, "xml"
        pdf_bytes = await convert_docx_to_pdf(content)
        if pdf_bytes:
            text = extract_text_from_pdf(pdf_bytes)
            if len(text) >= min_d:
                return text, "pdf_text"
            if len(text) < min_d:
                ocr = await extract_text_via_vision_qwen(pdf_bytes)
                if len(ocr.strip()) >= min_d:
                    return ocr.strip(), "vision_ocr"
        return text, "xml" if text else "empty"

    # Fallback: try python-docx then XML (e.g. odd extension but docx payload)
    text = extract_text_with_python_docx(content)
    if text and len(text) >= min_d:
        return text, "python_docx"
    text = extract_text_from_docx_xml(content)
    if text and len(text) >= min_d:
        return text, "xml"
    return text, "empty"
